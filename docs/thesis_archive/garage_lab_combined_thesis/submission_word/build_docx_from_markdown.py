from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches


MD_PATH = Path('garage_lab_combined/thesis/submission_word/THESIS_WORD_MASTER_75_85.md')
OUT_PATH = Path('garage_lab_combined/thesis/submission_word/THESIS_WORD_MASTER_75_85.docx')


def configure_document(doc: Document):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)


def format_paragraph(p):
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    # Keep heading readable but consistent
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h


def clean_inline(text: str) -> str:
    # basic markdown cleanup
    text = text.replace('**', '').replace('__', '')
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def add_table(doc: Document, rows):
    # rows: list[list[str]] where row 2 may be markdown separator
    # remove separator row like |---|---:|
    cleaned = []
    for r in rows:
        if all(re.fullmatch(r'[-: ]+', c or '') for c in r):
            continue
        cleaned.append(r)
    if not cleaned:
        return

    ncols = max(len(r) for r in cleaned)
    table = doc.add_table(rows=len(cleaned), cols=ncols)
    table.style = 'Table Grid'
    for i, r in enumerate(cleaned):
        for j in range(ncols):
            val = clean_inline(r[j]) if j < len(r) else ''
            cell = table.cell(i, j)
            cell.text = val
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.first_line_indent = Cm(0)


def parse_table_row(line: str):
    # line like | a | b | c |
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def build_doc(md_path: Path, out_path: Path):
    doc = Document()
    configure_document(doc)

    lines = md_path.read_text(encoding='utf-8').splitlines()
    i = 0
    base = md_path.parent

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            doc.add_paragraph('')
            i += 1
            continue

        # table block
        if s.startswith('|'):
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, block)
            continue

        # heading
        m = re.match(r'^(#{1,3})\s+(.*)$', s)
        if m:
            level = len(m.group(1))
            text = clean_inline(m.group(2).strip())
            # Word supports heading levels 1..9
            add_heading(doc, text, level)
            i += 1
            continue

        # image
        im = re.match(r'^!\[(.*?)\]\((.*?)\)$', s)
        if im:
            caption = clean_inline(im.group(1).strip())
            rel = im.group(2).strip()
            img = (base / rel).resolve()
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img), width=Inches(6.3))
                if caption:
                    cp = doc.add_paragraph(caption)
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.paragraph_format.line_spacing = 1.15
                    cp.paragraph_format.first_line_indent = Cm(0)
            else:
                p = doc.add_paragraph(f'[Missing image: {rel}]')
                format_paragraph(p)
            i += 1
            continue

        # bullet
        b = re.match(r'^-\s+(.*)$', s)
        if b:
            p = doc.add_paragraph(clean_inline(b.group(1).strip()), style='List Bullet')
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        # numbered
        n = re.match(r'^\d+\.\s+(.*)$', s)
        if n:
            p = doc.add_paragraph(clean_inline(n.group(1).strip()), style='List Number')
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        # default paragraph
        p = doc.add_paragraph(clean_inline(line))
        format_paragraph(p)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == '__main__':
    out = build_doc(MD_PATH, OUT_PATH)
    print(out)

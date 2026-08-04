#!/usr/bin/env python3
"""Build the Project_Cam English technical system report.

The script generates repository-native diagrams, converts the reviewed Markdown
source to DOCX with Pandoc, and applies a consistent Word layout.  It is fully
local and intentionally does not depend on external image assets or network
access.
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
import textwrap
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/reports/project_cam_technical_system_report_en.md"
ASSET_DIR = ROOT / "docs/assets/project_cam_report"
OUTPUT = ROOT / "docs/Project_Cam_Technical_System_Report_EN.docx"

FONT_REGULAR = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
FONT_BOLD = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
FONT_MONO = Path("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf")

NAVY = "102A43"
BLUE = "176B87"
TEAL = "00A6A6"
CYAN = "DDF6F4"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF3CD"
GOLD = "D69E2E"
RED = "C0392B"
PALE_RED = "FCE8E6"
GREEN = "2E7D5B"
PALE_GREEN = "E5F4EC"
INK = "243B53"
MUTED = "627D98"
LIGHT = "F5F8FA"
WHITE = "FFFFFF"
GRID = "C9D6E2"


def rgb(hex_value: str) -> tuple[int, int, int]:
    return tuple(int(hex_value[index : index + 2], 16) for index in (0, 2, 4))


def font(size: int, *, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_MONO if mono else FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(path), size)


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, face: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=face)[2] <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_text_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str = "",
    *,
    fill: str = WHITE,
    outline: str = BLUE,
    title_color: str = NAVY,
    body_color: str = INK,
    radius: int = 22,
    title_size: int = 30,
    body_size: int = 21,
    line_width: int = 4,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill), outline=rgb(outline), width=line_width)
    tf = font(title_size, bold=True)
    bf = font(body_size)
    inner_width = x2 - x1 - 36
    title_lines = wrapped_lines(draw, title, tf, inner_width)
    body_lines = wrapped_lines(draw, body, bf, inner_width) if body else []
    title_height = len(title_lines) * (title_size + 7)
    body_height = len(body_lines) * (body_size + 7)
    gap = 12 if body_lines else 0
    cursor = y1 + max(18, (y2 - y1 - title_height - body_height - gap) // 2)
    for line in title_lines:
        bbox = draw.textbbox((0, 0), line, font=tf)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, cursor), line, font=tf, fill=rgb(title_color))
        cursor += title_size + 7
    cursor += gap
    for line in body_lines:
        bbox = draw.textbbox((0, 0), line, font=bf)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, cursor), line, font=bf, fill=rgb(body_color))
        cursor += body_size + 7


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    color: str = MUTED,
    width: int = 5,
    head: int = 16,
) -> None:
    draw.line((start, end), fill=rgb(color), width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1.0)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    p1 = (x2 - head * ux + head * 0.65 * px, y2 - head * uy + head * 0.65 * py)
    p2 = (x2 - head * ux - head * 0.65 * px, y2 - head * uy - head * 0.65 * py)
    draw.polygon((end, p1, p2), fill=rgb(color))


def canvas(title: str, subtitle: str, *, height: int = 920) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, height), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 112), fill=rgb(NAVY))
    draw.text((58, 24), title, font=font(38, bold=True), fill=rgb(WHITE))
    draw.text((60, 72), subtitle, font=font(20), fill=rgb("CFE8F3"))
    return image, draw


def save_diagram(image: Image.Image, name: str) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image.save(ASSET_DIR / name, format="PNG", dpi=(240, 240), optimize=True)


def generate_system_architecture() -> None:
    image, draw = canvas(
        "Project_Cam system architecture",
        "Three planes share state, but actuation authority remains deliberately narrow.",
        height=980,
    )
    plane_y = [(145, 385), (405, 645), (665, 925)]
    plane_meta = [
        ("PERCEPTION PLANE", PALE_BLUE, BLUE),
        ("PRODUCT PLANE", CYAN, TEAL),
        ("ACTUATION & SAFETY PLANE", PALE_RED, RED),
    ]
    for (y1, y2), (label, fill, color) in zip(plane_y, plane_meta):
        draw.rounded_rectangle((36, y1, 1564, y2), radius=22, fill=rgb(fill), outline=rgb(color), width=3)
        draw.text((58, y1 + 16), label, font=font(22, bold=True), fill=rgb(color))

    boxes = [
        ((70, 215, 285, 340), "USB cameras", "latest frames"),
        ((335, 215, 575, 340), "2D inference", "pose + ball"),
        ((625, 215, 870, 340), "3D geometry", "association + SVD"),
        ((920, 215, 1160, 340), "Temporal state", "filters + tracks"),
        ((1210, 215, 1525, 340), "World snapshot", "quality + all people"),
        ((90, 475, 360, 600), "Live viewer", "2D / 3D overlays"),
        ((420, 475, 700, 600), "Assessment", "coaching screens"),
        ((760, 475, 1040, 600), "Training drills", "view-only states"),
        ((1100, 475, 1510, 600), "Desktop supervisor", "Tk committed · Tauri prototype"),
        ((90, 735, 375, 865), "Aim proposal", "bounded target state"),
        ((455, 735, 790, 865), "Fire-control boundary", "arm + fresh recheck"),
        ((870, 735, 1185, 865), "Fail-closed gates", "occupancy · age · epoch"),
        ((1265, 735, 1510, 865), "Launcher owner", "sole serial authority"),
    ]
    for index, (box, title, body) in enumerate(boxes):
        if index < 5:
            fill, outline = WHITE, BLUE
        elif index < 9:
            fill, outline = WHITE, TEAL
        else:
            fill, outline = WHITE, RED
        draw_text_box(draw, box, title, body, fill=fill, outline=outline, title_size=24, body_size=18)
    for x1, x2 in ((285, 335), (575, 625), (870, 920), (1160, 1210)):
        arrow(draw, (x1, 278), (x2, 278), color=BLUE)
    for x1, x2 in ((375, 455), (790, 870), (1185, 1265)):
        arrow(draw, (x1, 800), (x2, 800), color=RED)
    arrow(draw, (1365, 340), (1365, 475), color=TEAL)
    arrow(draw, (1330, 600), (1330, 735), color=RED)
    draw.text((985, 676), "Only structured aim/safety state crosses this boundary", font=font(19, bold=True), fill=rgb(RED))
    save_diagram(image, "system_architecture.png")


def generate_pose_geometry() -> None:
    image, draw = canvas(
        "From 2D keypoints to one world-space joint",
        "Coordinate conventions and observation timing are part of the estimate.",
        height=920,
    )
    draw_text_box(draw, (55, 160, 335, 295), "Source pixels", "distorted image coordinates", fill=PALE_BLUE)
    draw_text_box(draw, (390, 160, 690, 295), "Undistort", "K and lens coefficients", fill=PALE_BLUE)
    draw_text_box(draw, (745, 160, 1045, 295), "Normalized ray", "(x, y, 1) in camera frame", fill=PALE_BLUE)
    draw_text_box(draw, (1100, 160, 1545, 295), "Projection convention", "normalized points use P = [R | t]", fill=PALE_GOLD, outline=GOLD)
    for start, end in [((335, 228), (390, 228)), ((690, 228), (745, 228)), ((1045, 228), (1100, 228))]:
        arrow(draw, start, end, color=BLUE)

    camera_centres = [(245, 690), (520, 765), (1110, 770), (1370, 660)]
    point = (820, 485)
    for index, centre in enumerate(camera_centres, start=1):
        cx, cy = centre
        draw.rounded_rectangle((cx - 72, cy - 45, cx + 72, cy + 45), radius=14, fill=rgb(NAVY), outline=rgb(BLUE), width=3)
        label = f"Camera {index}"
        bbox = draw.textbbox((0, 0), label, font=font(20, bold=True))
        draw.text((cx - (bbox[2] - bbox[0]) / 2, cy - 13), label, font=font(20, bold=True), fill=rgb(WHITE))
        draw.line((centre, point), fill=rgb(TEAL), width=5)
        obs = (int(cx + (point[0] - cx) * 0.23), int(cy + (point[1] - cy) * 0.23))
        draw.ellipse((obs[0] - 10, obs[1] - 10, obs[0] + 10, obs[1] + 10), fill=rgb(GOLD))
    draw.ellipse((point[0] - 24, point[1] - 24, point[0] + 24, point[1] + 24), fill=rgb(RED), outline=rgb(NAVY), width=4)
    draw.text((point[0] + 38, point[1] - 24), "SVD estimate Xw", font=font(24, bold=True), fill=rgb(NAVY))
    draw.text((520, 340), "Stack  xi P3 - P1  and  yi P3 - P2  for every surviving camera", font=font(23, mono=True), fill=rgb(INK))
    warning = "Reprojection agreement does not prove synchronized capture, correct identity, or anatomical label."
    draw_text_box(
        draw,
        (460, 812, 1140, 895),
        warning,
        fill=PALE_RED,
        outline=RED,
        title_color=RED,
        title_size=18,
        line_width=3,
        radius=18,
    )
    save_diagram(image, "pose_geometry.png")


def generate_multi_person_flow() -> None:
    image, draw = canvas(
        "Multi-person tracking and identity flow",
        "Geometry creates the safety subject; Face ID only adds a local UX label.",
        height=940,
    )
    stages = [
        ((50, 190, 320, 330), "Per-view people", "boxes + 17 keypoints", PALE_BLUE, BLUE),
        ((375, 190, 675, 330), "Cross-view assignment", "candidate combinations", PALE_BLUE, BLUE),
        ((730, 190, 1010, 330), "3D persons", "pelvis + supported joints", PALE_BLUE, BLUE),
        ((1065, 190, 1545, 330), "Arena tracks", "monotonic track ID · primary epoch", PALE_GREEN, GREEN),
    ]
    for box, title, body, fill, outline in stages:
        draw_text_box(draw, box, title, body, fill=fill, outline=outline)
    for x1, x2 in ((320, 375), (675, 730), (1010, 1065)):
        arrow(draw, (x1, 260), (x2, 260), color=BLUE)

    draw_text_box(draw, (90, 470, 455, 640), "Round-robin face camera", "YuNet detect · SFace embed", fill=CYAN, outline=TEAL)
    draw_text_box(draw, (550, 470, 900, 640), "Voting gallery", "cosine threshold · repeated votes", fill=CYAN, outline=TEAL)
    draw_text_box(draw, (995, 470, 1495, 640), "Display-name attachment", "local coaching convenience · no liveness", fill=CYAN, outline=TEAL)
    arrow(draw, (455, 555), (550, 555), color=TEAL)
    arrow(draw, (900, 555), (995, 555), color=TEAL)
    arrow(draw, (1300, 470), (1300, 350), color=TEAL)

    draw_text_box(draw, (140, 745, 680, 875), "All-person safety snapshot", "geometry · support · age · primary track/epoch", fill=PALE_RED, outline=RED)
    draw_text_box(draw, (920, 745, 1460, 875), "Face identity excluded", "a name cannot authorize firing", fill=WHITE, outline=RED, title_color=RED)
    arrow(draw, (1250, 330), (600, 745), color=RED)
    draw.line((680, 810, 920, 810), fill=rgb(RED), width=6)
    draw.line((790, 770, 810, 850), fill=rgb(RED), width=8)
    draw.line((810, 770, 790, 850), fill=rgb(RED), width=8)
    save_diagram(image, "multi_person_flow.png")


def generate_fire_control() -> None:
    image, draw = canvas(
        "Fail-closed fire-control boundary",
        "The launcher owner rechecks fresh state immediately before one physical command.",
        height=930,
    )
    draw_text_box(draw, (55, 180, 315, 315), "Aim proposal", "target + mirror mode", fill=PALE_BLUE)
    draw_text_box(draw, (370, 180, 680, 315), "Arm context", "track ID + epoch + target", fill=PALE_GOLD, outline=GOLD)
    draw_text_box(draw, (735, 180, 1045, 315), "Fresh snapshot", "all people + support + age", fill=PALE_BLUE)
    draw_text_box(draw, (1100, 180, 1545, 315), "Immediate evaluation", "schema · identity epoch · occupancy", fill=PALE_RED, outline=RED)
    for start, end, color in [
        ((315, 248), (370, 248), BLUE),
        ((680, 248), (735, 248), GOLD),
        ((1045, 248), (1100, 248), RED),
    ]:
        arrow(draw, start, end, color=color)

    draw.rounded_rectangle((250, 410, 1350, 570), radius=34, fill=rgb(LIGHT), outline=rgb(NAVY), width=4)
    decision = "Are every required input, context, camera support, freshness, and corridor gate valid?"
    lines = wrapped_lines(draw, decision, font(30, bold=True), 1000)
    y = 444
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font(30, bold=True))
        draw.text((800 - (bbox[2] - bbox[0]) / 2, y), line, font=font(30, bold=True), fill=rgb(NAVY))
        y += 42

    draw_text_box(draw, (165, 690, 640, 850), "BLOCK + STOP", "missing · stale · changed · occupied", fill=PALE_RED, outline=RED, title_color=RED, title_size=34)
    draw_text_box(draw, (960, 690, 1435, 850), "ONE SHOOT COMMAND", "only the launcher owner writes serial", fill=PALE_GREEN, outline=GREEN, title_color=GREEN, title_size=31)
    arrow(draw, (550, 570), (400, 690), color=RED, width=7, head=22)
    arrow(draw, (1050, 570), (1200, 690), color=GREEN, width=7, head=22)
    draw.text((350, 610), "NO / ERROR", font=font(23, bold=True), fill=rgb(RED))
    draw.text((1120, 610), "YES", font=font(23, bold=True), fill=rgb(GREEN))
    draw.text((585, 882), "Implemented software boundary · not hardware commissioned or safety-certified", font=font(22, bold=True), fill=rgb(RED))
    save_diagram(image, "fire_control_boundary.png")


def generate_evidence_ladder() -> None:
    image, draw = canvas(
        "Evidence ladder",
        "Each level adds failure modes that lower levels cannot prove away.",
        height=940,
    )
    steps = [
        ("1", "Software contract", "unit/synthetic tests", 90, 790, 510, PALE_BLUE, BLUE),
        ("2", "Calibration fit", "intrinsic/extrinsic reprojection", 250, 655, 700, CYAN, TEAL),
        ("3", "Static ground truth", "accuracy + repeatability + failures", 430, 520, 900, PALE_GOLD, GOLD),
        ("4", "Dynamic validation", "motion + skew + occlusion + identity", 620, 385, 1110, PALE_GREEN, GREEN),
        ("5", "Commissioned operation", "hardware hazards + process + evidence", 820, 250, 1320, PALE_RED, RED),
    ]
    for number, title, body, x1, y1, x2, fill, outline in steps:
        y2 = y1 + 115
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill=rgb(fill), outline=rgb(outline), width=4)
        draw.ellipse((x1 + 18, y1 + 24, x1 + 82, y1 + 88), fill=rgb(outline))
        bbox = draw.textbbox((0, 0), number, font=font(30, bold=True))
        draw.text((x1 + 50 - (bbox[2] - bbox[0]) / 2, y1 + 37), number, font=font(30, bold=True), fill=rgb(WHITE))
        draw.text((x1 + 105, y1 + 20), title, font=font(27, bold=True), fill=rgb(NAVY))
        draw.text((x1 + 105, y1 + 66), body, font=font(20), fill=rgb(INK))
    draw_text_box(
        draw,
        (1080, 650, 1530, 860),
        "Current Project_Cam",
        "4-camera static evidence; 6-camera and product paths remain below commissioning",
        fill=WHITE,
        outline=RED,
        title_color=RED,
    )
    arrow(draw, (1080, 760), (930, 600), color=RED, width=6, head=20)
    save_diagram(image, "evidence_ladder.png")


def generate_diagrams() -> None:
    generate_system_architecture()
    generate_pose_geometry()
    generate_multi_person_flow()
    generate_fire_control()
    generate_evidence_ladder()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def select_table_style(style_names: set[str]) -> str:
    """Return the best table style available in a generated DOCX template."""
    if "Table Grid" in style_names:
        return "Table Grid"
    if "Table" in style_names:
        return "Table"
    raise ValueError("DOCX template does not provide a usable table style")


def inject_static_toc(markdown: str) -> str:
    """Insert a renderer-independent table of contents after YAML metadata."""
    headings = [
        match.group(1).strip()
        for line in markdown.splitlines()
        if (match := re.match(r"^# (?!#)(.+?)\s*$", line))
    ]
    toc_lines = ["# Table of Contents {.unnumbered}", ""]
    for index, heading in enumerate(headings, start=1):
        anchor = re.sub(r"[^\w\s-]", "", heading.lower(), flags=re.UNICODE)
        anchor = re.sub(r"[\s-]+", "-", anchor).strip("-")
        toc_lines.append(f"{index}. [{heading}](#{anchor})")
    toc = "\n".join(toc_lines) + "\n\n"

    lines = markdown.splitlines(keepends=True)
    if not lines or lines[0].strip() != "---":
        return toc + markdown
    for index, line in enumerate(lines[1:], start=1):
        if line.strip() == "---":
            return "".join(lines[: index + 1]) + "\n" + toc + "".join(lines[index + 1 :]).lstrip("\n")
    raise ValueError("unterminated YAML metadata block in report source")


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def set_run_font(run, name: str, size: float, color: str, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def style_document(input_docx: Path, output_docx: Path) -> None:
    doc = Document(input_docx)
    doc.core_properties.title = "Project_Cam: Multi-Camera Perception, Athlete Analytics, and Safe Launcher Research"
    doc.core_properties.subject = (
        "Technical system report, evidence snapshot 29 July 2026"
    )
    doc.core_properties.author = "Project_Cam Engineering"
    doc.core_properties.comments = "Generated from docs/reports/project_cam_technical_system_report_en.md"

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.widow_control = True

    style_specs = {
        "Title": ("Aptos Display", 29, NAVY, True, 0, 12),
        "Subtitle": ("Aptos", 15, TEAL, False, 0, 16),
        "Author": ("Aptos", 11, MUTED, False, 0, 4),
        "Date": ("Aptos", 10.5, MUTED, False, 0, 4),
        "Heading 1": ("Aptos Display", 17.5, NAVY, True, 15, 7),
        "Heading 2": ("Aptos Display", 13.2, BLUE, True, 11, 4),
        "Heading 3": ("Aptos", 11.2, TEAL, True, 8, 3),
        "TOC Heading": ("Aptos Display", 20, NAVY, True, 0, 12),
        "Caption": ("Aptos", 8.5, MUTED, False, 4, 8),
        "Block Text": ("Aptos", 10.25, NAVY, False, 6, 6),
    }
    for style_name, (name, size, color, bold, before, after) in style_specs.items():
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = name
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if style_name.startswith("Heading"):
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.keep_together = True

    if "Source Code" in styles:
        code = styles["Source Code"]
        code.font.name = "DejaVu Sans Mono"
        code.font.size = Pt(8)
        code.font.color.rgb = RGBColor.from_string(INK)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(5)

    title_seen = False
    toc_seen = False
    first_body_heading = False
    cover_break_added = False
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(86)
            paragraph.paragraph_format.space_after = Pt(15)
            title_seen = True
        elif style_name == "Subtitle":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif style_name in {"Author", "Date"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if style_name == "Date" and not cover_break_added:
                paragraph.add_run().add_break(WD_BREAK.PAGE)
                cover_break_added = True
        elif style_name == "TOC Heading":
            paragraph.paragraph_format.page_break_before = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            toc_seen = True
        elif style_name == "Heading 1":
            if paragraph.text.strip() == "Table of Contents":
                paragraph.paragraph_format.page_break_before = False
            elif not first_body_heading:
                paragraph.paragraph_format.page_break_before = True
                first_body_heading = True
        elif style_name == "Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = False
        if paragraph.text.startswith("Project_Cam is an engineering validation stack"):
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.right_indent = Cm(0.6)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(10)

    table_style = select_table_style({style.name for style in styles})
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = table_style
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                set_cell_shading(cell, NAVY if row_index == 0 else (LIGHT if row_index % 2 == 0 else WHITE))
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            "Aptos",
                            8.25,
                            WHITE if row_index == 0 else INK,
                            bold=True if row_index == 0 else None,
                        )

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.text = "PROJECT_CAM  /  TECHNICAL SYSTEM REPORT"
    for run in header_paragraph.runs:
        set_run_font(run, "Aptos", 8.2, NAVY, bold=True)
    p_pr = header_paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), TEAL)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_paragraph.add_run("Evidence snapshot 29 July 2026   ·   ")
    set_run_font(run, "Aptos", 8, MUTED)
    page_run = footer_paragraph.add_run()
    set_run_font(page_run, "Aptos", 8, NAVY, bold=True)
    add_field(page_run, "PAGE")

    # Keep images and their captions together where Word allows it.
    for paragraph in doc.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(7)
            paragraph.paragraph_format.space_after = Pt(3)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


IMAGE_LINK_RE = re.compile(r"!\[[^\]]*\]\(([^)\s]+)")


def referenced_figures(markdown: str) -> list[Path]:
    """Local figures the document links, resolved against the source dir."""
    found = []
    for target in IMAGE_LINK_RE.findall(markdown):
        if "://" in target:
            continue
        path = (SOURCE.parent / target.strip("<>")).resolve()
        if path not in found:
            found.append(path)
    return found


def verify_figures(markdown: str) -> list[Path]:
    """Fail loudly on a figure the document links but cannot resolve.

    Pandoc treats an unresolvable image as a warning and still exits 0, so
    ``check=True`` cannot catch it: the DOCX would ship to an external reader
    with the figure silently absent while the build printed success. The
    diagram filenames are written twice — once in ``generate_*`` and once in
    the Markdown link — so a rename on either side must be a build failure.
    """
    wanted = referenced_figures(markdown)
    missing = [p for p in wanted if not p.exists()]
    if missing:
        listing = "\n  ".join(str(p.relative_to(ROOT)) for p in missing)
        raise FileNotFoundError(
            f"report references {len(missing)} figure(s) that do not exist:\n  {listing}")
    return wanted


def build_docx() -> list[Path]:
    if not SOURCE.exists():
        raise FileNotFoundError(f"report source not found: {SOURCE}")
    if shutil.which("pandoc") is None:
        raise RuntimeError("pandoc is required to build the Word report")

    markdown = SOURCE.read_text(encoding="utf-8")
    figures = verify_figures(markdown)
    page_break = """```{=openxml}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```"""
    markdown = inject_static_toc(markdown)
    markdown = markdown.replace("<!-- PAGEBREAK -->", page_break)

    with tempfile.TemporaryDirectory(prefix="project_cam_report_") as temp_dir:
        temp = Path(temp_dir)
        prepared = temp / "report.md"
        intermediate = temp / "report.docx"
        prepared.write_text(markdown, encoding="utf-8")
        resource_path = f"{SOURCE.parent}:{ROOT / 'docs'}:{ROOT}"
        command = [
            "pandoc",
            str(prepared),
            "--from=markdown+raw_attribute+tex_math_dollars",
            "--to=docx",
            "--standalone",
            "--number-sections",
            f"--resource-path={resource_path}",
            "--metadata=link-citations:true",
            "-o",
            str(intermediate),
        ]
        result = subprocess.run(command, cwd=SOURCE.parent, check=True,
                                capture_output=True, text=True)
        # Second net: pandoc reports an unfetchable resource on stderr and
        # still exits 0, so a clean return code is not evidence of a complete
        # document.
        if "Could not fetch resource" in result.stderr:
            raise RuntimeError(
                "pandoc could not fetch a resource (it still exited 0):\n"
                + result.stderr.strip())
        if result.stderr.strip():
            print(result.stderr.strip())
        style_document(intermediate, OUTPUT)
    return figures


def main() -> None:
    generate_diagrams()
    figures = build_docx()
    print(f"Built {OUTPUT.relative_to(ROOT)}")
    used = {p.resolve() for p in figures}
    for path in sorted(ASSET_DIR.glob("*.png")):
        mark = "Generated" if path.resolve() in used else "Orphaned (unreferenced)"
        print(f"{mark} {path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
